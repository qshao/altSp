"""Integrated multi-source PCa alternative-splicing report (Word + Markdown).

Three parts mirror the prior reports (sequences -> function impact -> PCa/
resistance) but now span all sources, with a cross-source corroboration section.
"""
from __future__ import annotations
from datetime import date

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

MASTER = "results_collected/master_variants.tsv"
OUT_DOCX = "docs/Integrated_PCa_splicing_report.docx"
OUT_MD = "docs/Integrated_PCa_splicing_report.md"


def corroborated(df: pd.DataFrame) -> pd.DataFrame:
    d = df[df["corroborating_sources"].astype(int) >= 2].copy()
    return d.sort_values("corroborating_sources", ascending=False)


def _src_counts(df):
    return df["source"].value_counts().to_dict()


def build(master_tsv=MASTER, out_docx=OUT_DOCX, out_md=OUT_MD) -> dict:
    df = pd.read_csv(master_tsv, sep="\t").fillna("")
    df["corroborating_sources"] = df["corroborating_sources"].astype(int)
    multi = corroborated(df)
    multi_genes = sorted(multi["gene"].unique())
    counts = _src_counts(df)

    # ---------- Markdown ----------
    L = []
    L.append(f"# Integrated Prostate-Cancer Alternative-Splicing Report\n")
    L.append(f"_Multi-source: UniProt + TCGA SpliceSeq + literature + "
             f"ASCancerAtlas + GTEx baseline — {date.today():%Y-%m-%d}_\n")
    L.append("## Sources\n")
    for s, n in counts.items():
        L.append(f"- **{s}**: {n} variant rows")
    L.append(f"\nTotal: {len(df)} variant rows across {df['gene'].nunique()} genes.\n")
    L.append("## Cross-source corroboration (>=2 sources)\n")
    L.append("Genes supported by more than one independent source — the "
             "highest-confidence leads.\n")
    L.append("| Gene | Sources | Variants | Best PCa evidence |")
    L.append("|------|--------|----------|-------------------|")
    for g in multi_genes:
        sub = df[df["gene"] == g]
        srcs = ",".join(sorted(sub["source"].unique()))
        ev = next((x for x in sub["pca_evidence"] if x), "")
        L.append(f"| {g} | {srcs} | {len(sub)} | {ev} |")
    L.append("\n## Functional-impact highlights (domain / NMD / localization)\n")
    fi = df[(df["domains_hit"] != "") | (df["nmd_flag"] == "NMD-candidate")]
    for _, r in fi.sort_values("corroborating_sources", ascending=False).head(40).iterrows():
        L.append(f"- **{r['gene']}** ({r['source']}, {r['source_id']}): "
                 f"{r['change_class']}; domains hit: {r['domains_hit'] or 'n.a.'}; "
                 f"NMD: {r['nmd_flag']}; loc: {r['loc_flags'] or 'n.a.'}; "
                 f"GTEx: {r['gtex_baseline'] or 'n.a.'}")
    L.append("\n## Limitations\n")
    for s in [
        "Disease/resistance, NMD, and domain-disruption calls are predictions, "
        "not measurements.",
        "Cross-source corroboration is gene-level, not exon/variant-level.",
        "UniProt isoforms and literature variants are curated presence/absence; "
        "GTEx gives normal context only; cohort HRs are univariate single-cohort.",
        "CancerSplicingQTL was unreachable (HTTP 403) and is omitted.",
    ]:
        L.append(f"- {s}")
    open(out_md, "w").write("\n".join(L) + "\n")

    # ---------- Word ----------
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.add_heading("Integrated Prostate-Cancer Alternative-Splicing Report", 0)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"UniProt + TCGA + literature + ASCancerAtlas + GTEx — "
                f"{date.today():%Y-%m-%d}").italic = True

    doc.add_heading("Sources", 1)
    for s, n in counts.items():
        doc.add_paragraph(f"{s}: {n} variant rows", style="List Bullet")
    doc.add_paragraph(f"Total {len(df)} variant rows across "
                      f"{df['gene'].nunique()} genes.")

    doc.add_heading("Part 1 — Cross-source corroboration", 1)
    doc.add_paragraph("Genes supported by >=2 independent sources (highest-"
                      "confidence leads).")
    t = doc.add_table(rows=1, cols=4); t.style = "Light Grid Accent 1"
    for c, h in zip(t.rows[0].cells, ["Gene", "Sources", "Variants", "PCa evidence"]):
        c.paragraphs[0].add_run(h).bold = True
    for g in multi_genes:
        sub_df = df[df["gene"] == g]
        ev = next((x for x in sub_df["pca_evidence"] if x), "")
        cells = t.add_row().cells
        for c, v in zip(cells, [g, ",".join(sorted(sub_df["source"].unique())),
                                str(len(sub_df)), ev]):
            c.paragraphs[0].add_run(str(v)).font.size = Pt(8)

    doc.add_heading("Part 2 — Predicted functional impact", 1)
    doc.add_paragraph("Domain disruption, NMD candidacy, and localization "
                      "changes (UniProt-feature-driven; predictions).")
    fi2 = df[(df["domains_hit"] != "") | (df["nmd_flag"] == "NMD-candidate")]
    for _, r in fi2.sort_values("corroborating_sources", ascending=False).head(40).iterrows():
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{r['gene']} ").bold = True
        p.add_run(f"({r['source']}, {r['source_id']}): {r['change_class']}; "
                  f"domains: {r['domains_hit'] or 'n.a.'}; NMD: {r['nmd_flag']}; "
                  f"loc: {r['loc_flags'] or 'n.a.'}").font.size = Pt(9)

    doc.add_heading("Part 3 — Implications for PCa & therapeutic resistance", 1)
    doc.add_paragraph("AR-V7 (LBD-truncated, constitutively active) is the "
                      "headline resistance variant; PTEN/AR-axis and the "
                      "adhesion/invasion genes corroborated across sources are the "
                      "priority leads. All links are hypotheses.")

    doc.add_heading("Limitations", 1)
    for s in [
        "Predictions (NMD, domain disruption, disease links), not measurements.",
        "Corroboration is gene-level; exon-level matching is future work.",
        "GTEx is normal context; cohort HRs are univariate single-cohort.",
        "CancerSplicingQTL omitted (HTTP 403).",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    doc.save(out_docx)
    return {"rows": len(df), "genes": int(df["gene"].nunique()),
            "multi_source_genes": len(multi_genes)}


def main() -> None:
    stats = build()
    print(stats)


if __name__ == "__main__":
    main()
