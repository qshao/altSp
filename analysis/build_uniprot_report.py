"""Assemble the UniProt Word report (docs/UniProt_PCa_splicing_report.docx) with
the same three-part structure used for the TCGA report:

  Part 1 - before (canonical) / after (isoform) protein sequences
  Part 2 - predicted impact on each protein's known (UniProt-curated) function
  Part 3 - implications for prostate cancer and therapeutic resistance

Source: analysis/uniprot_sequence_impact.json  (run uniprot_sequence_impact.py).
Run:    PYTHONPATH=analysis python analysis/build_uniprot_report.py
"""
from __future__ import annotations
import json
from datetime import date

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from uniprot_report_notes import (
    resistance_note, TCGA_XREF, FEATURED_PCA_DISEASE, FEATURED_XREF,
)

IMPACT = "analysis/uniprot_sequence_impact.json"
OUT = "docs/UniProt_PCa_splicing_report.docx"
MONO = "Consolas"


def mono(p, text, size=7.5):
    run = p.add_run(text)
    run.font.name = MONO
    run.font.size = Pt(size)
    return run


def seq_block(doc, label, pid, seq, mark_pre, mark_suf):
    p = doc.add_paragraph()
    r = p.add_run(f"{label}  ({pid} | {len(seq)} aa)")
    r.bold = True
    r.font.size = Pt(9)
    body = doc.add_paragraph()
    body.paragraph_format.space_after = Pt(2)
    for i in range(0, len(seq), 60):
        mono(body, seq[i:i + 60] + "\n")
    notes = []
    if 0 < mark_pre < len(seq):
        notes.append(f"identical to the canonical protein up to residue "
                     f"{mark_pre}")
    if mark_suf > 0:
        notes.append(f"shares the final {mark_suf} residues")
    if notes:
        note = doc.add_paragraph()
        n = note.add_run("   ↑ " + "; ".join(notes) + ".")
        n.italic = True
        n.font.size = Pt(8)


def add_heading(doc, text, level):
    return doc.add_heading(text, level=level)


def pick_isoform(rec):
    """Choose the most informative isoform to feature for a gene.
    AR -> AR-V7; otherwise the largest-magnitude length change."""
    isos = rec["isoforms"]
    if rec["gene"] == "AR":
        v7 = [i for i in isos if "AR-V7" in i["synonyms"]]
        if v7:
            return v7[0]
    return max(isos, key=lambda i: abs(i["len_change"]) if i["len_change"]
               else (rec["canonical_len"] - i["identical_prefix"]))


def consequence_text(cls: str) -> str:
    if "truncation" in cls:
        return ("Loss of a terminal region is expected to remove the domains it "
                "encodes (loss-of-function, unless it yields a dominant fragment).")
    if cls == "substituted segment":
        return ("The native segment is replaced rather than simply removed, "
                "which can abolish a domain or create a neomorphic/constitutively "
                "active product (e.g. ligand-binding-domain loss).")
    if "deletion" in cls:
        return ("An in-frame internal deletion removes a cassette of residues; "
                "impact depends on whether the deleted segment carries a "
                "functional or regulatory element.")
    if "insertion" in cls:
        return ("An in-frame insertion adds residues that may extend or perturb "
                "a functional surface or linker.")
    return "Localised change; impact depends on overlap with functional regions."


def featured_block(doc, rec, part):
    iso = pick_isoform(rec)
    syn = f" ({'/'.join(iso['synonyms'])})" if iso["synonyms"] else ""
    title = f"{rec['gene']} — isoform {iso['isoform_name']}{syn}"
    add_heading(doc, title, 2)
    if part == 1:
        meta = doc.add_paragraph()
        meta.add_run(
            f"{rec['accession']} | canonical {rec['canonical_len']} aa "
            f"→ isoform {iso['after_len']} aa "
            f"({iso['len_change']:+d} aa)  |  change class: "
            f"{iso['change_class']}").font.size = Pt(9)
        if iso["varseq"]:
            v = doc.add_paragraph()
            v.add_run("UniProt VAR_SEQ: " + "; ".join(iso["varseq"])).font.size = Pt(8)
        seq_block(doc, "BEFORE (canonical)", rec["accession"],
                  rec["canonical_seq"], iso["identical_prefix"],
                  iso["identical_suffix"])
        seq_block(doc, f"AFTER (isoform {iso['isoform_name']})", iso["isoform_id"],
                  iso["after_seq"], iso["identical_prefix"], iso["identical_suffix"])
    elif part == 2:
        p = doc.add_paragraph()
        p.add_run("Known function (UniProt). ").bold = True
        p.add_run((rec["function"] or "Not curated.")[:700])
        p = doc.add_paragraph()
        p.add_run("Splice-induced change. ").bold = True
        p.add_run(iso["change_desc"])
        p = doc.add_paragraph()
        p.add_run("Predicted functional consequence. ").bold = True
        p.add_run(consequence_text(iso["change_class"]))
    elif part == 3:
        note = resistance_note(rec["gene"])
        x = TCGA_XREF.get(rec["gene"])
        p = doc.add_paragraph()
        if x:
            p.add_run(f"[TCGA cohort: {x['event']} ({x['type']}), "
                      f"{x['endpoint']} HR {x['HR']}, {x['direction']}, "
                      f"BH p {x['bh_p']:.1e}]  ").bold = True
        elif rec["is_prostate_disease"]:
            dis = [d["id"] for d in rec["diseases"]
                   if "prostate" in (d["id"] + d["desc"]).lower()]
            p.add_run(f"[UniProt disease: {'; '.join(dis)}]  ").bold = True
        p.add_run(note or "See UniProt disease annotation above.")


def main() -> None:
    data = json.load(open(IMPACT))
    by_gene = {d["gene"]: d for d in data}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    doc.add_heading(
        "Alternative Splicing of Prostate-Cancer Proteins (UniProt/Swiss-Prot):\n"
        "Curated Isoform Sequences, Functional Impact, and Links to Progression "
        "and Therapeutic Resistance", level=0)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"UniProtKB curated knowledge — companion to the TCGA "
                f"SpliceSeq report — {date.today():%Y-%m-%d}").italic = True

    # Executive summary
    add_heading(doc, "Executive summary", 1)
    n_pairs = sum(len(d["isoforms"]) for d in data)
    for s in [
        f"This report mirrors the TCGA-cohort analysis but starts from curated "
        f"knowledge: {len(data)} reviewed human proteins that UniProt annotates "
        f"as both alternatively spliced and prostate-cancer-associated, spanning "
        f"{n_pairs} canonical→isoform sequence pairs.",
        "For each protein we resolve the before (canonical) and after (isoform) "
        "protein sequences, classify the change, read across to UniProt's curated "
        "function and disease annotations, and reason about prostate-cancer "
        "progression and treatment resistance.",
        "The headline case is the androgen receptor: UniProt isoform 3 is AR-V7, "
        "the ligand-binding-domain-truncated, constitutively active variant that "
        "drives resistance to enzalutamide and abiraterone — the exact event "
        "the GRCh37-r75-based TCGA arm could not detect. This report therefore "
        "fills the principal blind spot of the cohort analysis.",
        "Twelve proteins carry a curated prostate-cancer disease annotation "
        "(including hereditary loci AR, ELAC2/HPC2, RNASEL/HPC1, MSR1, EHBP1, "
        "HNF1B, MSMB, plus PTEN, CHEK2, KLF6, MXI1, EPHB2); seven further genes "
        "independently surfaced as progression-associated protein-changing splices "
        "in the TCGA cohort (PRICKLE4, ITGA6, CTNND1, FES, STEAP3, SMN1, PRUNE2).",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    # Methods
    add_heading(doc, "Data and methods (brief)", 1)
    doc.add_paragraph(
        "Source: UniProtKB REST query (organism_id:9606) AND reviewed:true AND "
        "keyword:KW-0025 (Alternative splicing) AND \"prostate cancer\", retrieved "
        "into uniprot_prad_splicing/. For each protein the canonical sequence "
        "(\"before\") is paired with every annotated isoform sequence (\"after\") "
        "from isoforms.fasta. We compute the common prefix/suffix and classify the "
        "change (truncation, in-frame internal deletion/insertion, or substituted "
        "segment), and attach UniProt's VAR_SEQ feature descriptions, FUNCTION "
        "comment, and DISEASE annotations.")
    doc.add_paragraph(
        "Unlike the TCGA arm (statistical splice events mapped to GRCh37 r75 "
        "peptides, with univariate Cox hazard ratios), the isoforms here are "
        "expert-curated and include resistance variants such as AR-V7 that the "
        "genome build used in the cohort lacks. Where a gene also appears among "
        "the cohort's prioritized progression-associated events, the TCGA hazard "
        "ratio is shown for cross-validation. Disease links remain hypotheses for "
        "experimental follow-up.")

    featured = [g for g in FEATURED_PCA_DISEASE + FEATURED_XREF if g in by_gene]

    # Part 1
    add_heading(doc, "Part 1 — Before/after protein sequences", 1)
    doc.add_paragraph(
        "Canonical (\"before\") and the featured alternatively-spliced isoform "
        "(\"after\") are shown in full, with the residues at which they diverge "
        "from / re-converge on the canonical sequence. All isoform sequences are "
        "in uniprot_prad_splicing/isoforms.fasta; metrics for every pair are in "
        "analysis/uniprot_sequence_impact.json and the appendix.")
    for g in featured:
        featured_block(doc, by_gene[g], part=1)

    # Part 2
    add_heading(doc, "Part 2 — Predicted impact on protein function", 1)
    doc.add_paragraph(
        "Each protein's UniProt-curated function, the specific sequence change of "
        "the featured isoform, and the predicted consequence for that function.")
    for g in featured:
        featured_block(doc, by_gene[g], part=2)

    # Part 3
    add_heading(doc, "Part 3 — Implications for prostate cancer and "
                "therapeutic resistance", 1)
    doc.add_paragraph(
        "Linking the predicted functional change to prostate-cancer biology and "
        "treatment response. Bracketed tags give the curated UniProt disease term "
        "or, where available, the independent TCGA cohort hazard ratio.")
    for g in featured:
        featured_block(doc, by_gene[g], part=3)

    # Synthesis
    add_heading(doc, "Synthesis", 2)
    for s in [
        "Curated knowledge and cohort statistics are complementary: UniProt "
        "supplies the mechanistically characterised resistance isoforms (AR-V7, "
        "KLF6-SV1) that genome-build limitations hide from the TCGA arm, while the "
        "cohort supplies outcome associations for genes UniProt only lists "
        "qualitatively.",
        "Seven genes are supported by BOTH lines of evidence (UniProt PCa-splicing "
        "annotation AND a progression-associated protein-changing splice in the "
        "cohort), with internally consistent direction: adhesion/suppressor genes "
        "(PRUNE2, CTNND1, STEAP3) risk-down, invasion/proliferation genes (ITGA6, "
        "FES, PRICKLE4) risk-up.",
        "The androgen-receptor axis (AR/AR-V7, plus PTEN-driven AR-independent "
        "signalling) remains the dominant therapy-resistance theme and the highest "
        "priority for isoform-resolved follow-up (e.g. AR-V-aware RNA-seq).",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    # Limitations
    add_heading(doc, "Limitations", 1)
    for s in [
        "UniProt isoform catalogues are curated presence/absence, not "
        "quantitative: they do not say which isoform is expressed, or how much, in "
        "any given tumour.",
        "Functional consequences are inferred from sequence and curated domain "
        "knowledge, not measured here.",
        "Free-text \"prostate cancer\" matching can include proteins implicated "
        "indirectly; the 12 curated-disease genes are the high-confidence core.",
        "Cross-reference to the cohort is gene-level; matching the specific UniProt "
        "isoform to the specific TCGA splice event (exon-level) is the next step.",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    # Appendix
    add_heading(doc, "Appendix — all UniProt PCa splicing proteins", 1)
    doc.add_paragraph(
        f"All {len(data)} proteins with >=1 resolved isoform. \"Largest change\" "
        "summarises the isoform with the biggest length change. PCa = carries a "
        "curated prostate-cancer disease term; TCGA = also a prioritized cohort "
        "event.")
    cols = ["Gene", "Acc", "Canon aa", "#Iso", "Largest change", "Class",
            "PCa", "TCGA"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    for c, name in zip(table.rows[0].cells, cols):
        c.paragraphs[0].add_run(name).bold = True
        c.paragraphs[0].runs[0].font.size = Pt(8)
    for d in sorted(data, key=lambda x: (not x["is_prostate_disease"],
                                         x["gene"])):
        big = max(d["isoforms"], key=lambda i: abs(i["len_change"]))
        cells = table.add_row().cells
        vals = [d["gene"], d["accession"], d["canonical_len"],
                len(d["isoforms"]) + 1,
                f"{d['canonical_len']}→{big['after_len']} "
                f"({big['len_change']:+d})",
                big["change_class"],
                "yes" if d["is_prostate_disease"] else "",
                "yes" if d["gene"] in TCGA_XREF else ""]
        for c, v in zip(cells, vals):
            c.paragraphs[0].add_run(str(v)).font.size = Pt(7)

    doc.save(OUT)
    print(f"wrote {OUT}")
    print(f"featured genes: {featured}")
    print(f"appendix rows: {len(data)}")


if __name__ == "__main__":
    main()
