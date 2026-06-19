"""Assemble the Word report (docs/PRAD_splicing_report.docx) with three parts:

  Part 1 - before/after protein sequences for prioritized events
  Part 2 - predicted impact on each protein's known function
  Part 3 - implications for prostate cancer and therapeutic resistance

Run after sequence_impact.py.  python analysis/build_report.py
"""
from __future__ import annotations
import json
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from gene_annotations import get as annot

IMPACT = "analysis/sequence_impact.json"
OUT = "docs/PRAD_splicing_report.docx"

# Featured genes get full sequences + curated function/PCa text, in this order.
FEATURED = [
    "SPOP", "KLK2", "ACLY", "MAPT", "SMARCD3", "OGG1", "RBM6", "RAP1GAP",
    "VCL", "ADHFE1", "DOCK7", "SVIL", "SEC31A", "EXOC7", "HAUS5", "INO80E",
]
MONO = "Consolas"


def mono(p, text, size=7.5):
    run = p.add_run(text)
    run.font.name = MONO
    run.font.size = Pt(size)
    return run


def seq_block(doc, label, name, pid, seq, mark):
    """Print a sequence in 60-aa rows; annotate the divergence point."""
    p = doc.add_paragraph()
    r = p.add_run(f"{label}  ({name} | {pid} | {len(seq)} aa)")
    r.bold = True
    r.font.size = Pt(9)
    body = doc.add_paragraph()
    body.paragraph_format.space_after = Pt(2)
    for i in range(0, len(seq), 60):
        mono(body, seq[i:i + 60] + "\n")
    if 0 < mark < len(seq):
        note = doc.add_paragraph()
        n = note.add_run(f"   ↑ identical to the other isoform up to residue "
                         f"{mark}; sequence changes thereafter.")
        n.italic = True
        n.font.size = Pt(8)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h


def main() -> None:
    data = json.load(open(IMPACT))
    by_gene = {}
    for o in sorted(data, key=lambda x: x["bh_p"]):
        by_gene.setdefault(o["gene"], o)  # best event per gene

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ---- Title ----
    t = doc.add_heading(
        "Alternative Splicing in Prostate Adenocarcinoma:\n"
        "Protein-Sequence Consequences, Functional Impact, and Links to "
        "Progression and Therapeutic Resistance", level=0)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run(f"TCGA SpliceSeq (PRAD) — hypothesis-generating report — "
                    f"{date.today():%Y-%m-%d}")
    s.italic = True

    # ---- Executive summary ----
    add_heading(doc, "Executive summary", 1)
    summ = [
        f"Starting from {len(data)} differentially-spliced, protein-altering "
        "events that also associate with clinical outcome, this report resolves "
        "the before/after protein products, predicts how each splice changes the "
        "protein, and reasons about consequences for prostate-cancer progression "
        "and therapy resistance.",
        "Across the prioritized set, intron-retention and alternative-splice-site "
        "events overwhelmingly produce truncated or frame-shifted proteins (loss "
        "of C-terminal domains), and these losses skew toward worse "
        "progression-free outcome — a coherent loss-of-function signal.",
        "Featured events touch AR-pathway control (SPOP, KLK2), tumour "
        "lipid/energy metabolism (ACLY, ADHFE1, FADS3), DNA repair and chromatin "
        "(OGG1, SMARCD3, INO80E), the splicing machinery itself (RBM6), and the "
        "microtubule/taxane axis (MAPT) — each a plausible route to aggressive "
        "or treatment-resistant disease.",
    ]
    for s_ in summ:
        doc.add_paragraph(s_, style="List Bullet")

    # ---- Methods ----
    add_heading(doc, "Data and methods (brief)", 1)
    doc.add_paragraph(
        "Input: TCGA SpliceSeq PRAD table (67,534 events / 12,436 genes). Events "
        "were filtered for differential splicing (FDR < 0.05, |ΔPSI| ≥ 0.1) and "
        "their splice-in / splice-out isoforms mapped to Ensembl GRCh37 r75 "
        "transcripts and peptides. \"Before\" = representative splice-out protein; "
        "\"after\" = representative splice-in protein.")
    doc.add_paragraph(
        "Each event also carries univariate Cox hazard ratios (HR) against "
        "progression-free interval (PFI) and disease-free interval (DFI), with "
        "Benjamini–Hochberg-adjusted p-values. HR > 1 means high splice-in PSI "
        "tracks with worse outcome (\"risk-up\"). We retained protein-altering "
        "events with BH p < 0.05 and both before/after sequences resolved "
        f"({len(data)} events). Functional consequences are inferred from the "
        "sequence change; protein functions and prostate-cancer context are from "
        "the literature. All disease links are hypotheses for experimental "
        "follow-up, not established mechanisms.")

    # =====================================================================
    add_heading(doc, "Part 1 — Before/after protein sequences", 1)
    doc.add_paragraph(
        "For each featured gene the splice-out (“before”) and splice-in "
        "(“after”) protein sequences are shown in full, with the residue at "
        "which they diverge. Metrics for all prioritized events are in the "
        "appendix table; machine-readable sequences are in "
        "analysis/sequence_impact.json and results/proteins.fasta.")
    for g in FEATURED:
        if g not in by_gene:
            continue
        o = by_gene[g]
        add_heading(doc, f"{g} — {o['splice_event']} ({o['splice_type']})", 2)
        meta = doc.add_paragraph()
        meta.add_run(
            f"{o['before_len']} → {o['after_len']} aa ({o['net_aa_change']:+d} aa net)"
            f"  |  change class: {o['impact_class']}  |  splice class: "
            f"{o['comparison']}  |  {o['endpoint']} HR {o['HR']} "
            f"({o['direction']}), BH p {o['bh_p']:.1e}").font.size = Pt(9)
        seq_block(doc, "BEFORE (splice-out)", "out", o["before_pid"],
                  o["before_seq"], o["identical_prefix"])
        seq_block(doc, "AFTER (splice-in)", "in", o["after_pid"],
                  o["after_seq"], o["identical_prefix"])

    # =====================================================================
    add_heading(doc, "Part 2 — Predicted impact on protein function", 1)
    doc.add_paragraph(
        "For each gene: its established function, then how the observed "
        "before→after change is predicted to affect that function.")
    for g in FEATURED:
        if g not in by_gene:
            continue
        o = by_gene[g]
        fn, _ = annot(g)
        add_heading(doc, g, 2)
        p = doc.add_paragraph()
        p.add_run("Known function. ").bold = True
        p.add_run(fn)
        p = doc.add_paragraph()
        p.add_run("Splice-induced change. ").bold = True
        p.add_run(o["impact_detail"])
        p = doc.add_paragraph()
        p.add_run("Predicted functional consequence. ").bold = True
        if o["impact_class"] in ("truncation", "isoform loss (one side non-coding)"):
            p.add_run(
                "Substantial loss of the protein’s C-terminal region is expected "
                "to compromise or abolish its normal activity (loss-of-function).")
        elif o["impact_class"].startswith("frameshift"):
            p.add_run(
                "The C-terminal sequence is replaced rather than simply removed, "
                "potentially yielding a non-functional or neomorphic product.")
        else:
            p.add_run(
                "A localised in-frame change may modulate, rather than abolish, "
                "function depending on whether it overlaps a functional region.")

    # =====================================================================
    add_heading(doc, "Part 3 — Implications for prostate cancer and "
                "therapeutic resistance", 1)
    doc.add_paragraph(
        "Linking the predicted functional change (Part 2) to prostate-cancer "
        "biology and treatment response. Direction of the outcome association is "
        "noted; “risk-up” means the splice-in isoform tracks with worse "
        "progression.")
    for g in FEATURED:
        if g not in by_gene:
            continue
        o = by_gene[g]
        _, pca = annot(g)
        add_heading(doc, g, 2)
        p = doc.add_paragraph()
        p.add_run(f"[{o['endpoint']} HR {o['HR']}, {o['direction']}, "
                  f"BH p {o['bh_p']:.1e}]  ").bold = True
        p.add_run(pca)

    # ---- Synthesis ----
    add_heading(doc, "Synthesis", 2)
    for s_ in [
        "Loss-of-function splicing (intron retention, alternative splice sites) "
        "is the dominant harmful mode and is biased toward worse progression, "
        "suggesting splicing dysregulation erodes tumour-suppressor and "
        "genome-maintenance functions.",
        "Several featured genes converge on therapy-relevant axes: AR-pathway "
        "control (SPOP/KLK2), lipid/energy metabolism (ACLY/ADHFE1/FADS3), "
        "DNA-repair and chromatin remodelling (OGG1/SMARCD3/INO80E), and the "
        "microtubule/taxane interaction (MAPT).",
        "The androgen receptor itself shows alternate-terminator splicing with "
        "outcome signal but is not fully resolved by GRCh37 r75; re-analysis "
        "against an AR-V-aware annotation is the highest-priority follow-up.",
    ]:
        doc.add_paragraph(s_, style="List Bullet")

    # ---- Limitations ----
    add_heading(doc, "Limitations", 1)
    for s_ in [
        "Hazard ratios are univariate, single-cohort, and selected as the best of "
        "four endpoint/cut combinations — optimistic; use as ranking, not proof.",
        "Association is not causation; functional consequences are inferred from "
        "sequence, not measured.",
        "GRCh37 r75 misses key resistance variants (e.g. AR-V7), so the most "
        "therapy-relevant events are under-detected.",
        "Representative isoform per side simplifies multi-isoform events.",
    ]:
        doc.add_paragraph(s_, style="List Bullet")

    # ---- Appendix: full ranked table ----
    add_heading(doc, "Appendix — all prioritized protein-altering events", 1)
    doc.add_paragraph(
        f"All {len(data)} events with both protein sequences resolved and "
        "BH p < 0.05, ranked by significance.")
    cols = ["Gene", "Event", "Type", "Endpt", "HR", "Dir", "BH p",
            "Before→After", "Change"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    for c, name in zip(table.rows[0].cells, cols):
        c.paragraphs[0].add_run(name).bold = True
        c.paragraphs[0].runs[0].font.size = Pt(8)
    for o in sorted(data, key=lambda x: x["bh_p"]):
        cells = table.add_row().cells
        vals = [o["gene"], o["splice_event"], o["splice_type"], o["endpoint"],
                f"{o['HR']:g}", "up" if o["direction"] == "risk-up" else "down",
                f"{o['bh_p']:.1e}", f"{o['before_len']}→{o['after_len']}",
                o["impact_class"]]
        for c, v in zip(cells, vals):
            c.paragraphs[0].add_run(str(v)).font.size = Pt(7.5)

    doc.save(OUT)
    print(f"wrote {OUT}")
    print(f"featured genes present: "
          f"{[g for g in FEATURED if g in by_gene]}")
    print(f"appendix rows: {len(data)}")


if __name__ == "__main__":
    main()
